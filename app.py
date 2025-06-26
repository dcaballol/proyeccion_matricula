import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from prophet import Prophet
from io import BytesIO
from docx import Document
from docx.shared import Inches
import plotly.express as px
import zipfile
import tempfile
import os
import rarfile

st.set_page_config(layout="wide")
st.title("Proyección y Cobertura de Matrícula Escolar")
st.markdown("""
Panel para estimar y visualizar la proyección mensual y anual de matrícula por comuna, RBD y grado, 
cruzada con proyección demográfica, y calcular tasa de escolarización. Descarga Excel por comuna, por RBD y global.
""")

# ---- METODOLOGÍA (siempre visible para informe) ----
with st.expander("Metodología (haz clic para expandir)"):
    st.markdown("""
    **1. Proyección de matrícula:**  
    - Se utiliza Prophet para estimar matrícula mensual para el resto de 2025, 2026 y 2027, considerando tendencia y estacionalidad.
    - Se proyecta por cada combinación de comuna, RBD y grado.
    
    **2. Cruce demográfico:**  
    - Cada grado se asocia a una edad referencial (ej: 1° básico = 6 años, etc).
    - Se compara la matrícula proyectada (agregada por comuna, grado y año) con la población proyectada por comuna y edad.
    - Se calcula la tasa de escolarización: matrícula proyectada / población proyectada.
    
    **3. Visualización y exportación:**  
    - Gráficos mensuales, anuales y tasas.
    - Descarga de Excel con proyecciones y tasas, y generación de informes Word automáticos.
    """)

# ----------- CARGA DE DATOS ---------------
@st.cache_data
def cargar_matricula():
    rar_path = 'Asistencia 2015-2025.rar'
    csv_filename = 'Asistencia 2015-2025.csv'
    extract_dir = 'temp_extract'
    os.makedirs(extract_dir, exist_ok=True)
    extracted_csv = os.path.join(extract_dir, csv_filename)
    
    # Descomprime SOLO si no existe
    if not os.path.exists(extracted_csv):
        with rarfile.RarFile(rar_path) as rf:
            rf.extract(csv_filename, path=extract_dir)

    df = pd.read_csv(extracted_csv)
    df.columns = [col.lower().strip() for col in df.columns]
    df['mes'] = df['mes'].astype(int)
    df['agno'] = df['agno'].astype(int)
    df['cod_grado'] = df['cod_grado'].astype(str)
    df['fecha'] = pd.to_datetime(df['agno'].astype(str) + '-' + df['mes'].astype(str) + '-01')
    return df
@st.cache_data
def cargar_poblacion():
    df = pd.read_excel('poblacion.xlsx')
    df.columns = [col.lower().strip() for col in df.columns]
    df['agno'] = df['agno'].astype(int)
    df['edad'] = df['edad'].astype(int)
    return df

df_mat = cargar_matricula()
df_demo = cargar_poblacion()

map_grado_edad = {
    'PK': 4, 'K': 5,
    '1BAS': 6, '2BAS': 7, '3BAS': 8, '4BAS': 9, '5BAS': 10,
    '6BAS': 11, '7BAS': 12, '8BAS': 13,
    '1MED': 14, '2MED': 15, '3MED': 16, '4MED': 17,
}

df_mat['edad'] = df_mat['cod_grado'].map(map_grado_edad)

# ----------- SELECCIÓN DINÁMICA ---------------
comunas = sorted(df_mat['comuna'].dropna().unique())
comunas_sel = st.multiselect("Selecciona comuna(s)", comunas, default=comunas[:1])
rbd_all = df_mat[df_mat['comuna'].isin(comunas_sel)]['rbd'].dropna().unique()
rbds_sel = st.multiselect("Selecciona RBD(s)", sorted(rbd_all), default=rbd_all[:1])
grados_all = df_mat[df_mat['rbd'].isin(rbds_sel)]['cod_grado'].dropna().unique()
grados_sel = st.multiselect("Selecciona Grado(s)", sorted(grados_all), default=grados_all[:1])

df_sel = df_mat[
    (df_mat['comuna'].isin(comunas_sel)) &
    (df_mat['rbd'].isin(rbds_sel)) &
    (df_mat['cod_grado'].isin(grados_sel))
].copy()
df_sel = df_sel.sort_values(['comuna','rbd','cod_grado','fecha'])

combos = df_sel[['comuna','rbd','cod_grado','edad']].drop_duplicates().values.tolist()

st.info(f"{len(combos)} combinaciones seleccionadas.")

# ----------- PROYECCIÓN Y VISUALIZACIÓN ---------------

results = []
for comuna, rbd, cod_grado, edad in combos:
    grupo = df_sel[(df_sel['comuna']==comuna)&(df_sel['rbd']==rbd)&(df_sel['cod_grado']==cod_grado)].copy()
    if len(grupo) < 24: continue
    ts = grupo[['fecha','matricula']].rename(columns={'fecha':'ds','matricula':'y'})
    ts = ts.set_index('ds').asfreq('MS').reset_index()
    m = Prophet(yearly_seasonality=True, seasonality_mode='multiplicative')
    m.fit(ts)
    last_date = ts['ds'].max()
    periods = (12-last_date.month+1) + 12*2  # resto 2025 + 2026 + 2027
    future = m.make_future_dataframe(periods=periods, freq='MS')
    forecast = m.predict(future)
    forecast['mes'] = forecast['ds'].dt.month
    forecast['agno'] = forecast['ds'].dt.year
    forecast['comuna'] = comuna
    forecast['rbd'] = rbd
    forecast['cod_grado'] = cod_grado
    forecast['edad'] = edad
    # Toma solo meses >= actual si es 2025
    current_year = pd.Timestamp.today().year
    forecast = forecast[(forecast['agno'] >= current_year) & (forecast['agno'] <= 2027)]
    results.append({'comuna':comuna, 'rbd':rbd, 'cod_grado':cod_grado, 'edad':edad,
                    'hist': ts, 'forecast': forecast})

# ----------- VISUALIZACIÓN Y AGREGADOS ---------------

# a) Matrícula mensual proyectada y anual total
all_forecast = pd.concat([res['forecast'] for res in results]) if results else pd.DataFrame()
all_forecast['yhat'] = all_forecast['yhat'].clip(lower=0)
# Agregados anuales
anual = (
    all_forecast.groupby(['comuna','rbd','cod_grado','agno','edad'])['yhat']
    .sum().reset_index().rename(columns={'yhat':'matricula_proyectada'})
)

# b) Cruce con demografía y tasa escolarización
df_demo_sub = df_demo[(df_demo['comuna'].isin(comunas_sel)) & (df_demo['agno'].isin([2025,2026,2027]))]
anual = anual.merge(df_demo_sub, on=['comuna','agno','edad'], how='left')
anual['tasa_escolarizacion'] = anual['matricula_proyectada'] / anual['poblacion']

# --------- GRÁFICOS AGREGADOS --------
st.subheader("Totales anuales y tasas proyectadas")

if not anual.empty:
    fig1 = px.bar(anual, x='agno', y='matricula_proyectada', color='cod_grado',
                  barmode='group', facet_row='comuna', facet_col='rbd',
                  title="Matrícula proyectada anual por grado, RBD y comuna")
    st.plotly_chart(fig1, use_container_width=True)

    fig2 = px.line(anual, x='agno', y='tasa_escolarizacion', color='cod_grado',
                   facet_row='comuna', facet_col='rbd',
                   title="Tasa de escolarización proyectada por año")
    st.plotly_chart(fig2, use_container_width=True)

# --------- EXPORTACIÓN EXCEL POR COMUNA ---------
for comuna in anual['comuna'].unique():
    df_com = anual[anual['comuna']==comuna]
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_com.to_excel(writer, sheet_name=f"Totales_{comuna}", index=False)
    st.download_button(
        label=f"Descargar Excel de Totales para {comuna}",
        data=output.getvalue(),
        file_name=f"Totales_Anuales_{comuna}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# --------- EXPORTACIÓN EXCEL POR RBD ---------
for rbd in anual['rbd'].unique():
    df_rbd = anual[anual['rbd']==rbd]
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_rbd.to_excel(writer, sheet_name=f"Totales_{rbd}", index=False)
    st.download_button(
        label=f"Descargar Excel de Totales para RBD {rbd}",
        data=output.getvalue(),
        file_name=f"Totales_Anuales_RBD_{rbd}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

# --------- EXPORTACIÓN EXCEL GLOBAL ---------
output = BytesIO()
with pd.ExcelWriter(output, engine='openpyxl') as writer:
    anual.to_excel(writer, sheet_name="Totales_Anuales_Global", index=False)
st.download_button(
    label="Descargar Excel de Totales - Global",
    data=output.getvalue(),
    file_name="Totales_Anuales_Global.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

# --------- DESCARGA INFORME WORD (POR COMBINACIÓN) ---------
def genera_word(res, anual):
    doc = Document()
    doc.add_heading("Informe de Proyección de Matrícula Escolar y Tasa de Escolarización", 0)
    doc.add_heading("Metodología", level=1)
    doc.add_paragraph(
        "Se utiliza Prophet para proyección mensual de matrícula por establecimiento, grado y comuna para el resto de 2025, y años completos 2026 y 2027. "
        "Cada grado se asocia a un tramo etario y se calcula la tasa de escolarización al cruzar la matrícula proyectada con la proyección demográfica comunal."
    )
    doc.add_heading("Resultados", level=1)
    doc.add_paragraph(f"Comuna: {res['comuna']} - RBD: {res['rbd']} - Grado: {res['cod_grado']} ({res['edad']} años)")
    # Gráfico mensual
    fig, ax = plt.subplots(figsize=(7,3))
    ax.plot(res['hist']['ds'], res['hist']['y'], marker='o', label='Histórico')
    ax.plot(res['forecast']['ds'], res['forecast']['yhat'], color='orange', label='Proyección')
    ax.fill_between(res['forecast']['ds'], res['forecast']['yhat_lower'], res['forecast']['yhat_upper'], color='orange', alpha=0.2)
    ax.set_title("Matrícula mensual (histórico y proyección)")
    ax.set_xlabel("Fecha"); ax.set_ylabel("Matrícula")
    ax.legend()
    img_buf = BytesIO()
    fig.savefig(img_buf, format='png'); plt.close(fig)
    img_buf.seek(0)
    doc.add_picture(img_buf, width=Inches(5.5))
    # Tabla anual y tasa
    tabla = anual[
        (anual['comuna']==res['comuna'])&
        (anual['rbd']==res['rbd'])&
        (anual['cod_grado']==res['cod_grado'])
    ][['agno','matricula_proyectada','poblacion','tasa_escolarizacion']]
    doc.add_heading("Totales anuales y tasa escolarización", level=2)
    t = doc.add_table(rows=1, cols=tabla.shape[1])
    for i, col in enumerate(tabla.columns):
        t.cell(0,i).text = str(col)
    for row in tabla.itertuples(index=False):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = f"{val:.2f}" if isinstance(val, float) else str(val)
    return doc

if st.button("Descargar informes Word (ZIP)"):
    with tempfile.TemporaryDirectory() as tempdir:
        zipbuf = BytesIO()
        with zipfile.ZipFile(zipbuf, 'w') as zipf:
            for res in results:
                doc = genera_word(res, anual)
                fname = f"Informe_{res['comuna']}_{res['rbd']}_{res['cod_grado']}.docx"
                file_path = os.path.join(tempdir, fname)
                doc.save(file_path)
                zipf.write(file_path, arcname=fname)
        zipbuf.seek(0)
        st.download_button(
            label="Descargar ZIP con todos los informes Word",
            data=zipbuf,
            file_name="Informes_Matricula_y_Tasa.zip",
            mime="application/zip"
        )

st.success("¡Panel listo para Streamlit Cloud, con exportaciones por comuna, por RBD y global!")

